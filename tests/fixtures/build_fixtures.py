"""Synthesize the three golden control packages, deterministically.

C23024 - quarterly partner rebate calculation (numeric recompute + sign-off)
C10032 - monthly consolidation reconciliation (zero-delta + GMT/CDT temporal
         + email->ZIP->workbook nesting + IC-elims scope exclusion)
C10075 - quarterly EMR review (screenshot OCR tie-out + SoD sign-off)

Everything is fixed content and fixed dates: rebuilding produces the same
facts, so ingest hashes are stable run over run.
"""
from __future__ import annotations

import io
import zipfile
from datetime import datetime, timezone, timedelta
from email.message import EmailMessage
from email.utils import format_datetime
from pathlib import Path

import docx as docx_mod
import openpyxl
from PIL import Image, ImageDraw, ImageFont

FIXTURES = Path(__file__).parent / "controls"


# ------------------------------------------------------------------ helpers

def _wb(sheets: dict[str, dict[str, object]]) -> bytes:
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    for name, cells in sheets.items():
        ws = wb.create_sheet(name)
        for coord, value in cells.items():
            ws[coord] = value
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _eml(sender: str, to: str, subject: str, date: datetime, body_lines: list[str],
         message_id: str, attachments: list[tuple[str, bytes]] | None = None) -> bytes:
    msg = EmailMessage()
    msg["From"] = sender
    msg["To"] = to
    msg["Subject"] = subject
    msg["Date"] = format_datetime(date)
    msg["Message-ID"] = message_id
    msg.set_content("\n".join(body_lines))
    for filename, data in (attachments or []):
        msg.add_attachment(data, maintype="application", subtype="octet-stream",
                           filename=filename)
    return msg.as_bytes()


def _zip(files: list[tuple[str, bytes]]) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in files:
            info = zipfile.ZipInfo(name, date_time=(2026, 6, 3, 8, 0, 0))
            zf.writestr(info, data)
    return buf.getvalue()


def _screenshot(lines: list[str]) -> bytes:
    font = None
    for cand in ("/System/Library/Fonts/Supplemental/Arial.ttf",
                 "/System/Library/Fonts/Helvetica.ttc",
                 "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"):
        try:
            font = ImageFont.truetype(cand, 42)
            break
        except OSError:
            continue
    img = Image.new("RGB", (900, 80 + 90 * len(lines)), "white")
    d = ImageDraw.Draw(img)
    y = 40
    for line in lines:
        d.text((40, y), line, fill="black", font=font)
        y += 90
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _doc404(path: Path, paragraphs: list[str]) -> None:
    d = docx_mod.Document()
    for p in paragraphs:
        d.add_paragraph(p)
    d.save(str(path))


def _docx_simple(paragraphs: list[str]) -> bytes:
    d = docx_mod.Document()
    for p in paragraphs:
        d.add_paragraph(p)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# ------------------------------------------------------------------ controls

def build_C23024(root: Path) -> None:
    pkg = root / "C23024" / "package"
    pkg.mkdir(parents=True, exist_ok=True)
    (pkg / "rebate_Q2_2026.xlsx").write_bytes(_wb({
        "Sales": {"A1": "Region", "B1": "Q2 sales",
                  "A2": "AMS", "B2": 125000.0, "A3": "EMEA", "B3": 98000.0,
                  "A4": "APJ", "B4": 143000.0, "A5": "LATAM", "B5": 88000.0,
                  "A7": "Total", "B7": 454000.0},
        "Rebate": {"A2": "Total sales", "B2": 454000.0,
                   "A3": "Rebate rate", "B3": 0.02,
                   "A4": "Rebate payable", "B4": 9080.0},
        "Meta": {"A1": "Prepared", "B1": "2026-07-05 14:00 GMT",
                 "A2": "Preparer", "B2": "priya.sharma@example.com"}}))
    (pkg / "approval_C23024.eml").write_bytes(_eml(
        "ravi.mehta@example.com", "priya.sharma@example.com",
        "RE: Q2 FY26 partner rebate calculation - approval",
        datetime(2026, 7, 6, 9, 30, tzinfo=timezone(timedelta(hours=-5))),
        ["Priya,", "",
         "Reviewed and approved the Q2 partner rebate calculation.",
         "Rebate payable of 9,080.00 ties to the certified sales extract.", "",
         "Ravi Mehta", "Controller, Print Finance"],
        "<c23024-approval-q2fy26@example.com>"))
    _doc404(root / "C23024" / "404_C23024.docx", [
        "Control C23024 - Quarterly Partner Rebate Calculation. The financial "
        "analyst recomputes partner rebates from the certified sales extract; "
        "the controller reviews and approves the calculation before posting.",
        "EXPECTED EVIDENCE:",
        "- e_wb | Rebate calculation workbook | hints: rebate | required",
        "- e_approval | Controller approval email | hints: approval_C23024 | required",
        "CHECKS:",
        '- n1 | numeric | Regional sales foot to the quarterly total | '
        '{"op":"sum_equals","source":{"evidence":"e_wb","sheet":"Sales","cells":["B2","B3","B4","B5"]},'
        '"target":{"evidence":"e_wb","sheet":"Sales","cell":"B7"},"tolerance":0.01}',
        '- n2 | numeric | Rebate payable equals total sales times rate | '
        '{"op":"product_equals","source":{"evidence":"e_wb","sheet":"Rebate","cells":["B2","B3"]},'
        '"target":{"evidence":"e_wb","sheet":"Rebate","cell":"B4"},"tolerance":0.01}',
        '- s1 | signoff | Controller approval postdates preparation, SoD held | '
        '{"approval_email_evidence":"e_approval","preparer":"priya.sharma",'
        '"prepared_at":{"evidence":"e_wb","sheet":"Meta","cell":"B1","tz":"GMT"}}',
        "SIGNOFF: preparer=Financial Analyst; approver=Controller; distinct=true; order=true"])


def build_C10032(root: Path) -> None:
    pkg = root / "C10032" / "package"
    pkg.mkdir(parents=True, exist_ok=True)
    (pkg / "consolidation_recon_May2026.xlsx").write_bytes(_wb({
        "Recon": {"A2": "TB per certified WEBI", "B2": 5432100.25,
                  "A3": "TB per consolidation ledger", "B3": 5432100.25,
                  "A5": "Delta", "B5": 0.0},
        "Meta": {"A1": "WEBI run (GMT)", "B1": "2026-06-03 09:15",
                 "A2": "Preparer", "B2": "lena.fischer@example.com"}}))
    support_zip = _zip([
        ("recon_checklist.docx", _docx_simple(
            ["Consolidation reconciliation checklist - May 2026",
             "TB tie-out performed; IC eliminations monitored under control 10033."])),
        ("ic_elims_summary.xlsx", _wb({
            "Elims": {"A2": "IC eliminations May", "B2": 78400.0,
                      "A3": "Cleared", "B3": 61200.0,
                      "A5": "Open (monitored separately)", "B5": 17200.0}}))])
    (pkg / "approval_C10032.eml").write_bytes(_eml(
        "tomas.alvarez@example.com", "lena.fischer@example.com",
        "RE: May 2026 consolidation reconciliation - approval",
        datetime(2026, 6, 3, 8, 45, tzinfo=timezone(timedelta(hours=-5))),  # CDT
        ["Lena,", "",
         "Approved - the May consolidation reconciliation ties to the certified WEBI TB.",
         "Support attached for the file.", "",
         "Tomas Alvarez", "Consolidations Manager"],
        "<c10032-approval-may26@example.com>",
        attachments=[("recon_support.zip", support_zip)]))
    _doc404(root / "C10032" / "404_C10032.docx", [
        "Control C10032 - Monthly Consolidation Reconciliation. The consolidation "
        "analyst reconciles the certified WEBI trial balance to the consolidation "
        "ledger; the consolidations manager approves after the WEBI run. "
        "Intercompany eliminations are monitored separately under control 10033.",
        "EXPECTED EVIDENCE:",
        "- e_wb | Consolidation reconciliation workbook | hints: consolidation_recon | required",
        "- e_approval | Manager approval email with support | hints: approval_C10032 | required",
        "CHECKS:",
        '- d1 | numeric | WEBI TB to ledger delta is zero | '
        '{"op":"delta_zero","source":{"evidence":"e_wb","sheet":"Recon","cells":["B2","B3"]},'
        '"target":{"evidence":"e_wb","sheet":"Recon","cell":"B5"},"tolerance":0.01}',
        '- t1 | temporal | Approval postdates the certified WEBI run (GMT vs CDT) | '
        '{"earlier":{"evidence":"e_wb","sheet":"Meta","cell":"B1","tz":"GMT"},'
        '"later_email_evidence":"e_approval"}',
        "SCOPE EXCLUSIONS:",
        "- x1 | IC eliminations monitored separately under control 10033 | hints: ic_elims",
        "SIGNOFF: preparer=Consolidation Analyst; approver=Consolidations Manager; distinct=true; order=true"])


def build_C10075(root: Path) -> None:
    pkg = root / "C10075" / "package"
    pkg.mkdir(parents=True, exist_ok=True)
    (pkg / "emr_workbook_Q2_2026.xlsx").write_bytes(_wb({
        "EMR": {"A7": "Total Qty per EMR extract", "C7": 4820,
                "A8": "Reporting period", "C8": "Q2 FY26"},
        "Meta": {"A1": "Prepared", "B1": "2026-07-10 11:00 GMT",
                 "A2": "Preparer", "B2": "dana.wu@example.com"}}))
    (pkg / "screenshot_qty.png").write_bytes(_screenshot(
        ["EMR Extract - Q2 FY26", "Total Qty: 4,820", "Source: certified WEBI folder"]))
    (pkg / "approval_C10075.eml").write_bytes(_eml(
        "marco.rossi@example.com", "dana.wu@example.com",
        "RE: Q2 FY26 EMR review - sign-off",
        datetime(2026, 7, 10, 16, 20, tzinfo=timezone.utc),
        ["Dana,", "",
         "Reviewed and approved the Q2 EMR quarterly review.",
         "Screenshot quantities tie to the workbook extract.", "",
         "Marco Rossi", "External Reporting Manager"],
        "<c10075-approval-q2fy26@example.com>"))
    _doc404(root / "C10075" / "404_C10075.docx", [
        "Control C10075 - Quarterly External Manufacturing Report (EMR) Review. "
        "The analyst ties screenshot IPE from the certified WEBI folder to the "
        "EMR workbook; the external reporting manager signs off. Preparer and "
        "reviewer must be distinct.",
        "EXPECTED EVIDENCE:",
        "- e_wb | EMR review workbook | hints: emr_workbook | required",
        "- e_shot | Screenshot IPE of certified extract | hints: screenshot | required",
        "- e_approval | Manager sign-off email | hints: approval_C10075 | required",
        "CHECKS:",
        '- v1 | vision | Screenshot total quantity ties to workbook | '
        '{"label":"Total Qty","image_evidence":"e_shot",'
        '"target":{"evidence":"e_wb","sheet":"EMR","cell":"C7"},"tolerance":0}',
        '- s1 | signoff | Manager sign-off postdates preparation, SoD held | '
        '{"approval_email_evidence":"e_approval","preparer":"dana.wu",'
        '"prepared_at":{"evidence":"e_wb","sheet":"Meta","cell":"B1","tz":"GMT"}}',
        "SIGNOFF: preparer=Reporting Analyst; approver=External Reporting Manager; distinct=true; order=true"])


CONTROLS = {"C23024": ("quarterly", build_C23024),
            "C10032": ("monthly", build_C10032),
            "C10075": ("quarterly", build_C10075)}


def build_all(root: Path = FIXTURES) -> Path:
    for cid, (_freq, fn) in CONTROLS.items():
        fn(root)
    return root


if __name__ == "__main__":
    print(build_all())
